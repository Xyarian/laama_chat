"""
Laama Chat - A Streamlit app for chatting with AI LLM models using RAG
"""
import os
import streamlit as st
from htmlTemplates import get_full_css
from llama_index.core import VectorStoreIndex, Document, Settings
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.llms.ollama import Ollama
import ollama
import chardet
import magic
import pdfplumber
from docx import Document as DocxDocument
from chats_db import create_database, save_chat, load_chats, load_chat_messages, delete_chat, get_default_model, set_default_model
import logging

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("laama_chat")

# Mapping of model display names to internal model names
MODEL_MAPPING = {
    "Llama 3.1": "llama3.1",
    "Gemma 2": "gemma2",
}

# Set environment variables for offline mode
os.environ['HF_DATASETS_OFFLINE'] = '1'
os.environ['TRANSFORMERS_OFFLINE'] = '1'
os.environ['HF_HUB_OFFLINE'] = '1'

# Specify a local cache directory for the model
local_model_path = st.secrets.local_model.path # online usage "BAAI/bge-base-en-v1.5"
cache_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), "../.model_cache"))
os.environ['TRANSFORMERS_CACHE'] = cache_dir
os.environ['HF_HOME'] = cache_dir

@st.cache_resource
def initialize_llamaindex():
    """Initialize LlamaIndex components."""
    try:
        # Attempt to load the model in offline mode
        embed_model = HuggingFaceEmbedding(
            model_name=local_model_path, # "BAAI/bge-base-en-v1.5"
            trust_remote_code=False,
            cache_folder=cache_dir,
            local_files_only=True  # This ensures only local files are used
        )
        logger.info(f"Model loaded from local cache: {cache_dir}")
    except Exception as e:
        logger.error(f"Failed to load model in offline mode: {e}")
        logger.error("Please download the model files to the specified cache directory before running in offline mode.")
        st.error("Failed to load the AI model in offline mode. Please check the logs for more information.")
        st.stop()

    Settings.embed_model = embed_model
    return VectorStoreIndex([])

def get_ollama_llm(model_name):
    """Get Ollama LLM based on the model name."""
    return Ollama(model=model_name, request_timeout=360.0)

def clear_index():
    """Clear the index and reset the RAG flag."""
    st.session_state.index = initialize_llamaindex()
    st.session_state.use_rag = False

def on_file_change():
    """Callback function when file uploader state changes."""
    if 'uploaded_file' in st.session_state:
        if st.session_state.uploaded_file is None:
            # File was removed
            clear_index()
            st.toast("File removed. RAG disabled.", icon="🗑️")
    else:
        # First time upload, do nothing
        pass

def get_ai_response(messages, model):
    """Get a response from the AI model based on the user messages."""
    try:
        response = ollama.chat(
            model=model,
            messages=messages,
        )
        return response['message']['content']
    except Exception as e:
        logger.error(f"Laama response error: {e}")
        return "I'm sorry, I couldn't process your request."

def detect_encoding(file_content):
    """Detect the encoding of the file content."""
    result = chardet.detect(file_content)
    return result['encoding']

def is_text_file(file_content):
    """Check if the file content is likely to be text."""
    mime = magic.Magic(mime=True)
    file_type = mime.from_buffer(file_content)
    return file_type.startswith('text/') or 'charset' in file_type

def extract_text(file):
    """Extract text content from a file based on its content."""
    try:
        file_content = file.read()
        mime = magic.Magic(mime=True)
        file_type = mime.from_buffer(file_content)
        
        if is_text_file(file_content):
            encoding = detect_encoding(file_content)
            return file_content.decode(encoding), file_type
        
        # Handle PDF files
        elif file.type == "application/pdf":
            return extract_text_from_pdf(file), file_type
        
        # Handle DOCX files
        elif file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return extract_text_from_docx(file), file_type
        
        else:
            st.warning(f"Unsupported file type: {file.type}. Treating as plain text.")
            encoding = detect_encoding(file_content)
            return file_content.decode(encoding), file_type
    
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return ""

def extract_text_from_pdf(file):
    """Extract text from a PDF file using pdfplumber."""
    try:
        text = ''
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text
                else:
                    st.warning("Some pages may not contain extractable text.")
        return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return ""

def extract_text_from_docx(file):
    """Extract text from a DOCX file using python-docx."""
    try:
        doc = DocxDocument(file)
        text = '\n'.join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return ""

def delete_chat_callback(chat_id, chat_name):
    """Callback function to delete a chat by ID."""
    try:
        delete_chat(chat_id)
        st.session_state.saved_chats = load_chats()
        st.toast(f"Successfully deleted chat ***{chat_name}***.", icon="🗑️")
    except Exception as e:
        st.error(f"Error deleting chat: {e}")

@st.cache_resource
def initialize_database():
    try:
        create_database()
        logger.info("Database initialized successfully")
    except Exception as e:
        logger.error(f"Error initializing database: {e}")

def main():
    """Main function for the Laama Chat app."""
    # Set page config and title
    st.set_page_config(
        page_title="Laama Chat",
        page_icon="🦙",
        layout="centered",
        initial_sidebar_state="collapsed",
        menu_items={
            'About': "# Laama Chat RAG v0.1.0-alpha  \nBy Xyarian 2024  \nhttps://github.com/Xyarian"
        }
    )
    
    # Initialize index and RAG flag
    if 'index' not in st.session_state:
        st.session_state.index = initialize_llamaindex()
    if 'use_rag' not in st.session_state:
        st.session_state.use_rag = False
    # Initialize database
    initialize_database()
    
    # Load and apply custom CSS
    image_file = "files/bg_image.png"
    image_ext = "png"
    full_css = get_full_css(image_file, image_ext)
    st.write(full_css, unsafe_allow_html=True)
    
    st.title('Laama Chat')

    # Initialize session state
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'saved_chats' not in st.session_state:
        st.session_state.saved_chats = load_chats()

    # Get the default model
    default_model = get_default_model()
    default_model_display_name = next((name for name, value in MODEL_MAPPING.items() if value == default_model), "Llama 3.1")

    # Sidebar setup
    st.sidebar.header("Laama Chat Controls", help="Manage your current chat session", divider=True)
    
    col1, col2 = st.sidebar.columns(2)
    with col1:
        selected_model_display_name = st.selectbox(
            "Select AI Model",
            list(MODEL_MAPPING.keys()),
            index=list(MODEL_MAPPING.keys()).index(default_model_display_name),
            label_visibility="collapsed"
        )
        selected_model = MODEL_MAPPING[selected_model_display_name]
        if selected_model != default_model:
            set_default_model(selected_model)
        
    with col2:    
        new_chat_button = st.button("🆕 New Chat", help="Start a new chat session")

    # Update LlamaIndex LLM based on selected model
    Settings.llm = get_ollama_llm(selected_model)

    st.write(f"Welcome to Laama Chat with {selected_model_display_name}! Ask me anything or share a document to get started.")

    # Handle new chat button
    if new_chat_button:
        st.session_state.messages = []  # Clear the chat messages
        clear_index()  # Clear the index and reset RAG flag
        st.toast("You can start a new conversation.", icon="✅")
        st.toast("Chat session cleared and RAG disabled.", icon="🗑️")

    # Save Chat Section
    with st.sidebar.form("save_chat_form"):
        chat_name = st.text_input("Enter a chat name for saving:", placeholder="My Chat", max_chars=50)
        submit_button = st.form_submit_button("💾 Save Chat", help="Save the current chat with a custom name")

    if submit_button and chat_name:
        try:
            save_chat(chat_name, st.session_state.messages, selected_model)
            st.sidebar.success("Chat saved successfully!")
            st.toast(f"Chat ***{chat_name}*** saved successfully.", icon="✅")
            st.session_state.saved_chats = load_chats()
        except Exception as e:
            st.sidebar.error(f"Error saving chat: {e}")

    # Saved Chats Section
    st.sidebar.subheader("Saved Chats", help="Load or delete saved chats", divider=True)
    for chat_id, chat_name, chat_model in st.session_state.saved_chats:
        col1, col2 = st.sidebar.columns([3, 1])
        with col1:
            if st.button(f"{chat_name} ({chat_model})", help="Load the saved chat", key=f"load_{chat_id}", use_container_width=True):
                try:
                    success, messages, model = load_chat_messages(chat_id)
                    if success:
                        st.session_state.messages = messages
                        selected_model = model
                        selected_model_display_name = next((name for name, value in MODEL_MAPPING.items() if value == model), "Llama 3.1")
                        Settings.llm = get_ollama_llm(selected_model)
                        if messages:
                            st.toast(f"Successfully loaded messages for chat ***{chat_name}*** using model {selected_model_display_name}.", icon="✅")
                        else:
                            st.toast(f"Chat ***{chat_name}*** has been loaded but contains no messages.", icon="ℹ️")
                except Exception as e:
                    st.error(f"Error loading chat messages: {e}")
        with col2:
            st.button("🗑️", help="Delete the saved chat", key=f"delete_{chat_id}", on_click=delete_chat_callback, args=(chat_id, chat_name))

    # File uploader for adding files to index
    uploaded_file = st.file_uploader("Upload a file", type=None,  # Allow all file types
                                    on_change=on_file_change, key="uploaded_file")
    
    if uploaded_file is not None:
        if not st.session_state.use_rag:  # Only process if RAG is not already enabled
            file_content, file_type = extract_text(uploaded_file)
            if file_content:
                document = Document(text=file_content)
                st.session_state.index.insert(document)
                st.session_state.use_rag = True  # Enable RAG when a document is added
                st.success(f"File `{uploaded_file.name} ({file_type})` content added to the index. RAG enabled.")
            else:
                st.warning("The file appears to be empty or unreadable.")

    # Display chat messages
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])

    # Handle user input
    if prompt := st.chat_input("Message Laama Chat..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)

        with st.spinner("Laama is thinking..."):
            if st.session_state.use_rag:
                # Use LlamaIndex to query the index
                query_engine = st.session_state.index.as_query_engine()
                response = query_engine.query(prompt)
                ai_response = str(response)
            else:
                # Use regular Ollama chat API
                ai_response = get_ai_response(st.session_state.messages, selected_model)

        with st.chat_message("assistant"):
            st.markdown(ai_response)
        st.session_state.messages.append({"role": "assistant", "content": ai_response})

if __name__ == "__main__":
    main()
